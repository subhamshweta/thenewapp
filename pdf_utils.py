# pdf_utils.py
import PyPDF2
import io
from fpdf import FPDF
import tempfile
import textwrap
from docx import Document
import os
import re
import logging

# Set up logging
logging.basicConfig(filename='resume_enhancer.log', level=logging.DEBUG, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

def extract_text_from_document(file):
    """
    Extract text content from a PDF or DOCX file.
    
    Args:
        file: File object of the PDF or DOCX
        
    Returns:
        str: Extracted text from the document
    """
    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension == '.pdf':
            return extract_text_from_pdf(file)
        elif file_extension == '.docx':
            return extract_text_from_docx(file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        logging.error(f"Error extracting text from document: {e}")
        return "Error extracting text from document"

def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() or ""
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return "Error extracting text from PDF"

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return "Error extracting text from DOCX"

def create_document(resume_text, output_format='pdf'):
    if output_format.lower() == 'pdf':
        return create_pdf(resume_text)
    elif output_format.lower() == 'docx':
        return create_docx(resume_text)
    else:
        raise ValueError(f"Unsupported output format: {output_format}")

def create_pdf(resume_text):
    """
    Create a professionally formatted PDF document with a modern two-column layout.
    Falls back to basic formatting if parsing fails.
    
    Args:
        resume_text (str): The optimized resume text
        
    Returns:
        bytes: PDF file as bytes
    """
    try:
        # Log the input resume text for debugging
        logging.debug(f"Input resume text:\n{resume_text}")

        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_margins(15, 15, 15)
        sidebar_width = 50
        main_content_x = 70
        main_content_width = pdf.w - pdf.r_margin - main_content_x

        sections = parse_markdown_resume(resume_text)

        # Check if parsing was successful
        required_sections = ['name', 'contact', 'summary', 'skills', 'experience', 'education']
        missing_sections = [s for s in required_sections if s not in sections or not sections[s].strip()]
        if missing_sections:
            logging.warning(f"Missing or empty sections: {missing_sections}")
            return create_fallback_pdf(resume_text)

        # Sidebar: Contact and Skills
        pdf.set_xy(15, 15)
        if 'contact' in sections:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(sidebar_width, 8, "CONTACT", ln=True)
            pdf.set_font("Arial", size=9)
            contact_lines = sections['contact'].split('\n')
            contact_text = ' | '.join(line.strip() for line in contact_lines if line.strip())
            pdf.multi_cell(sidebar_width, 5, contact_text)
            pdf.ln(5)

        if 'skills' in sections:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(sidebar_width, 8, "SKILLS", ln=True)
            pdf.set_font("Arial", size=9)
            skills = sections['skills'].split('\n')
            for skill in skills:
                if skill.strip():
                    pdf.multi_cell(sidebar_width, 5, f"• {skill.strip()}")
            pdf.ln(5)

        # Main Content: Name, Summary, Experience, Education, etc.
        pdf.set_xy(main_content_x, 15)
        if 'name' in sections:
            pdf.set_font("Arial", 'B', 18)
            pdf.cell(0, 12, sections['name'].upper(), ln=True)
            pdf.ln(3)

        pdf.line(main_content_x, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

        if 'summary' in sections:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "PROFESSIONAL SUMMARY", ln=True)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(main_content_width, 5, sections['summary'])
            pdf.ln(5)

        if 'experience' in sections:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "PROFESSIONAL EXPERIENCE", ln=True)
            jobs = sections['experience'].split('\n\n')
            for job in jobs:
                if job.strip():
                    job_parts = job.split('\n', 1)
                    pdf.set_font("Arial", 'B', 10)
                    pdf.cell(0, 5, job_parts[0].strip(), ln=True)
                    if len(job_parts) > 1:
                        pdf.set_font("Arial", size=10)
                        details = job_parts[1].strip()
                        for line in details.split('\n'):
                            if line.strip():
                                pdf.cell(5, 5, '', ln=0)
                                pdf.cell(5, 5, '•', ln=0)
                                pdf.multi_cell(main_content_width - 10, 5, line.strip()[1:] if line.strip().startswith('•') else line.strip())
                    pdf.ln(3)

        if 'education' in sections:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, "EDUCATION", ln=True)
            edu_entries = sections['education'].split('\n\n')
            for entry in edu_entries:
                if entry.strip():
                    entry_parts = entry.split('\n', 1)
                    pdf.set_font("Arial", 'B', 10)
                    pdf.cell(0, 5, entry_parts[0].strip(), ln=True)
                    if len(entry_parts) > 1:
                        pdf.set_font("Arial", size=10)
                        details = entry_parts[1].strip()
                        for line in details.split('\n'):
                            if line.strip():
                                pdf.multi_cell(main_content_width, 5, line.strip())
                    pdf.ln(3)

        for section_name in ['certifications', 'projects', 'awards', 'publications', 'hobbies_interests']:
            if section_name in sections:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 8, section_name.upper(), ln=True)
                pdf.set_font("Arial", size=10)
                pdf.multi_cell(main_content_width, 5, sections[section_name])
                pdf.ln(5)

        pdf_output = pdf.output(dest='S')
        return pdf_output if isinstance(pdf_output, bytes) else pdf_output.encode('latin-1')
    except Exception as e:
        logging.error(f"Error creating PDF: {e}")
        return create_fallback_pdf(resume_text)

def parse_markdown_resume(markdown_text):
    """
    Parse a markdown-formatted resume into sections with flexible header matching.
    
    Args:
        markdown_text (str): The markdown-formatted resume
    Returns:
        dict: Dictionary of sections
    """
    sections = {}
    current_section = None
    current_content = []
    lines = markdown_text.split('\n')
    section_patterns = [
        ('name', r'NAME'),
        ('contact', r'CONTACT'),
        ('summary', r'PROFESSIONAL SUMMARY|SUMMARY'),
        ('skills', r'SKILLS'),
        ('experience', r'PROFESSIONAL EXPERIENCE|EXPERIENCE|WORK EXPERIENCE'),
        ('education', r'EDUCATION'),
        ('certifications', r'CERTIFICATIONS'),
        ('projects', r'PROJECTS'),
        ('hobbies_interests', r'HOBBIES\s*&\s*INTERESTS|HOBBIES|INTERESTS'),
        ('awards', r'AWARDS'),
        ('publications', r'PUBLICATIONS')
    ]

    for line in lines:
        line = line.strip()
        matched = False
        for section_key, pattern in section_patterns:
            # Match headers with flexible markdown levels and colons
            if re.match(rf'^(#+)?\s*{pattern}[\s:]*$', line, re.IGNORECASE):
                if current_section:
                    sections[current_section] = '\n'.join(current_content).strip()
                current_section = section_key
                current_content = []
                matched = True
                break
        if not matched:
            current_content.append(line)
    
    if current_section:
        sections[current_section] = '\n'.join(current_content).strip()

    # Ensure all required sections exist
    required_sections = ['name', 'contact', 'summary', 'skills', 'experience', 'education']
    for section in required_sections:
        if section not in sections or not sections[section].strip():
            sections[section] = f"[Missing {section.capitalize()} Section]"
            logging.warning(f"{section.capitalize()} section not found or empty in parsed resume")

    logging.debug(f"Parsed sections: {sections.keys()}")
    return sections

def create_fallback_pdf(resume_text):
    """
    Create a basic PDF with raw resume text if parsing fails.
    
    Args:
        resume_text (str): The optimized resume text
        
    Returns:
        bytes: PDF file as bytes
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_margins(15, 15, 15)
        pdf.set_font("Arial", size=10)
        pdf.multi_cell(0, 5, "Note: Structured formatting could not be applied due to parsing issues.\nBelow is the raw optimized resume text:\n")
        pdf.ln(5)
        lines = resume_text.split('\n')
        for line in lines:
            if line.strip():
                pdf.multi_cell(0, 5, line.strip())
        pdf_output = pdf.output(dest='S')
        return pdf_output if isinstance(pdf_output, bytes) else pdf_output.encode('latin-1')
    except Exception as e:
        logging.error(f"Error creating fallback PDF: {e}")
        return create_error_document('PDF')

def create_docx(resume_text):
    try:
        doc = Document()
        sections = parse_markdown_resume(resume_text)
        
        for section_name, content in sections.items():
            if section_name == 'name':
                doc.add_heading(content.upper(), level=1)
            else:
                doc.add_heading(section_name.upper(), level=2)
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip(), style='ListBullet' if line.strip().startswith('•') else None)
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue()
    except Exception as e:
        logging.error(f"Error creating DOCX: {e}")
        return create_error_document('DOCX')

def create_error_document(format_type):
    try:
        if format_type.upper() == 'PDF':
            error_pdf = FPDF()
            error_pdf.add_page()
            error_pdf.set_font("Arial", size=12)
            error_pdf.cell(0, 10, "Error creating optimized resume PDF", ln=True)
            error_pdf.cell(0, 10, "Possible issues: Invalid markdown format or missing sections.", ln=True)
            error_pdf.cell(0, 10, "Please check the resume text and try again.", ln=True)
            error_output = error_pdf.output(dest='S')
            return error_output if isinstance(error_output, bytes) else error_output.encode('latin-1')
        elif format_type.upper() == 'DOCX':
            error_doc = Document()
            error_doc.add_heading("Error creating optimized resume DOCX", 0)
            error_doc.add_paragraph("Possible issues: Invalid markdown format or missing sections.")
            error_doc.add_paragraph("Please check the resume text and try again.")
            error_bytes = io.BytesIO()
            error_doc.save(error_bytes)
            error_bytes.seek(0)
            return error_bytes.getvalue()
    except Exception as e:
        logging.error(f"Error creating error document: {e}")
        return b""
